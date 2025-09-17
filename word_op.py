import docx
import re


def extract_questions(doc):
    questions = []
    pattern = re.compile(r'^(\d+)\.\[([^\]]+)\[(.*)')
    current_question = None

    for para in doc.paragraphs:
        text = para.text.strip()
        match = pattern.match(text)
        if match:
            # If a new question is found, save the current question if it exists
            if current_question:
                questions.append(current_question)
            question_number = match.group(1)
            question_type = match.group(2)
            question_content = match.group(3).strip()
            current_question = (question_number, question_type, question_content)

        elif current_question:
            # Append additional lines to the current question content
            current_question = (current_question[0], current_question[1], current_question[2] + " " + text)

    # Append the last question if it exists
    if current_question:
        questions.append(current_question)

    return questions

from pypinyin import pinyin, lazy_pinyin
def sort_by_pinyin(questions):
    def pinyin_key(question):
        return lazy_pinyin(question[2][0])[0] if question[2] else ""

    return sorted(questions, key=pinyin_key)


def save_sorted_questions(questions, output_path):
    new_doc = docx.Document()

    for question in questions:
        new_doc.add_paragraph(f"{question[0]}.[{question[1]}] {question[2]}")

    new_doc.save(output_path)
def sort_by_tihao(questions):
    def pinyin_key(question):
        return lazy_pinyin(question[1][0])[0] if question[1] else ""

    return sorted(questions, key=pinyin_key)
from collections import OrderedDict
def remove_duplicates(questions):
    unique_questions = list(OrderedDict.fromkeys(questions))
    return unique_questions

def main(input_path, output_path=None):
    doc = docx.Document(input_path)
    questions = extract_questions(doc)
    sorted_questions = sort_by_tihao(questions)
    sorted_questions=sort_by_pinyin(sorted_questions)

    save_sorted_questions(sorted_questions,output_path)



